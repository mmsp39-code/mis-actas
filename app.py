import streamlit as st
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import io
import os

# Configuración para Pantalla Ancha
st.set_page_config(page_title="Generador de Actas Profesional", layout="wide")

# --- FUNCIÓN DE REINICIO TOTAL ---
def restart_application():
    # Borra todas las claves guardadas en la sesión (textos y archivos)
    for key in st.session_state.keys():
        del st.session_state[key]
    # Forzar recarga de la página
    st.rerun()

# --- BARRA LATERAL ---
with st.sidebar:
    st.header("📋 Datos del Acta")
    # Añadimos 'key' a cada input para que el reinicio los detecte
    edar = st.text_input("EDAR", placeholder="Nombre de la planta", key="edar_name")
    idcoste = st.text_input("IDCOSTE", placeholder="Referencia IDCOSTE", key="id_coste")
    poblacion = st.text_input("Población", key="pob")
    direccion = st.text_input("Dirección", key="dir")
    provincia = st.text_input("Provincia", key="prov")
    fecha = st.text_input("Fecha instalación", placeholder="DD/MM/AAAA", key="fec")
    tecnicos = st.text_input("Técnicos instaladores", key="tec")
    responsable = st.text_input("Responsable Explotación", key="resp")
    
    st.divider()
    # BOTÓN DE REINICIAR (Limpia textos y fotos)
    if st.button("♻️ REINICIAR TODO", use_container_width=True, type="secondary"):
        restart_application()

# --- PANEL PRINCIPAL ---
st.title("📄 Generador de Actas ADASA & INELCOM")
st.divider()

# --- CARGA DE FOTOS EN 4 COLUMNAS ---
st.write("### 📸 Gestión de Archivos por Categoría")
c1, c2, c3, c4 = st.columns(4)

with c1:
    st.error("#### 🖼️ GENERALES")
    f_portada = st.file_uploader("Foto Portada", accept_multiple_files=True, key="u1")
    f_cartel = st.file_uploader("Cartel Informativo", accept_multiple_files=True, key="u2")
    f_graficas = st.file_uploader("Gráficas", accept_multiple_files=True, key="u3")

with c2:
    st.warning("#### 🌊 ALIVIOS")
    f_alivio_e1 = st.file_uploader("Alivio EDAR 1", accept_multiple_files=True, key="u4")
    f_alivio_e2 = st.file_uploader("Alivio EDAR 2", accept_multiple_files=True, key="u5")
    f_alivio_c1 = st.file_uploader("Alivio Colector 1", accept_multiple_files=True, key="u6")
    f_alivio_c2 = st.file_uploader("Alivio Colector 2", accept_multiple_files=True, key="u7")
    f_alivio_c3 = st.file_uploader("Alivio Colector 3", accept_multiple_files=True, key="u8")

with c3:
    st.success("#### 📥 CAUDALÍMETROS")
    f_cauda_e1 = st.file_uploader("Caudalímetro Entrada 1", accept_multiple_files=True, key="u9")
    f_cauda_e2 = st.file_uploader("Caudalímetro Entrada 2", accept_multiple_files=True, key="u10")
    f_cauda_s1 = st.file_uploader("Caudalímetro Salida 1", accept_multiple_files=True, key="u11")
    f_cauda_s2 = st.file_uploader("Caudalímetro Salida 2", accept_multiple_files=True, key="u12")

with c4:
    st.info("#### 🧪 CALIDAD")
    f_calid_e1 = st.file_uploader("Calidad Entrada 1", accept_multiple_files=True, key="u13")
    f_calid_e2 = st.file_uploader("Calidad Entrada 2", accept_multiple_files=True, key="u14")
    f_calid_s1 = st.file_uploader("Calidad Salida 1", accept_multiple_files=True, key="u15")
    f_calid_s2 = st.file_uploader("Calidad Salida 2", accept_multiple_files=True, key="u16")

st.divider()

if st.button("🚀 GENERAR DOCUMENTO FINAL", use_container_width=True, type="primary"):
    if not edar:
        st.error("⚠️ El nombre de la EDAR es obligatorio para generar el documento.")
    else:
        with st.spinner("Maquetando acta..."):
            doc = Document()
            
            # 1. Portada y Datos
            if os.path.exists('logo_instituciona.png'):
                p_inst = doc.add_paragraph()
                p_inst.alignment = 1
                p_inst.add_run().add_picture('logo_instituciona.png', width=Inches(5))

            doc.add_heading(edar, 0).alignment = 1
            doc.add_heading('ACTA DE CERTIFICACIÓN', 1).alignment = 1

            if f_portada:
                p_port = doc.add_paragraph()
                p_port.alignment = 1
                p_port.add_run().add_picture(f_portada[0], width=Inches(4))

            # Tabla de Datos
            datos = [("EDAR", edar), ("IDCOSTE", idcoste), ("Población", poblacion), ("Dirección", direccion), ("Provincia", provincia), ("Fecha instalación", fecha), ("Técnicos instaladores", tecnicos), ("Responsable Explotación", responsable)]
            tbl = doc.add_table(rows=len(datos), cols=2)
            tbl.style = 'Table Grid'
            for i, (k, v) in enumerate(datos):
                tbl.rows[i].cells[0].text, tbl.rows[i].cells[1].text = k, str(v)

            # Logos Portada
            p_logos = doc.add_paragraph()
            p_logos.alignment = 1
            if os.path.exists('logo_adasa.png'): p_logos.add_run().add_picture('logo_adasa.png', width=Inches(1.2))
            p_logos.add_run("      ") 
            if os.path.exists('logo_inelcom.png'): p_logos.add_run().add_picture('logo_inelcom.png', width=Inches(1.2))

            # 2. Secciones
            leyenda_eq = "Fotos instalación y ubicación de equipos y sondas instalados."
            conclusion_fija = "La instalación en EDAR, queda completada correctamente y en servicio."

            secciones = [
                ("CARTEL INFORMATIVO", f_cartel, "Fotografía del cartel de subvenciones."),
                ("ALIVIO EDAR 1", f_alivio_e1, leyenda_eq), ("ALIVIO EDAR 2", f_alivio_e2, leyenda_eq),
                ("ALIVIO COLECTOR 1", f_alivio_c1, leyenda_eq), ("ALIVIO COLECTOR 2", f_alivio_c2, leyenda_eq), ("ALIVIO COLECTOR 3", f_alivio_c3, leyenda_eq),
                ("CAUDALÍMETRO ENTRADA 1", f_cauda_e1, leyenda_eq), ("CAUDALÍMETRO ENTRADA 2", f_cauda_e2, leyenda_eq),
                ("CAUDALÍMETRO SALIDA 1", f_cauda_s1, leyenda_eq), ("CAUDALÍMETRO SALIDA 2", f_cauda_s2, leyenda_eq),
                ("CALIDAD ENTRADA 1", f_calid_e1, leyenda_eq), ("CALIDAD ENTRADA 2", f_calid_e2, leyenda_eq),
                ("CALIDAD SALIDA 1", f_calid_s1, leyenda_eq), ("CALIDAD SALIDA 2", f_calid_s2, leyenda_eq),
                ("GRÁFICAS", f_graficas, conclusion_fija)
            ]

            for titulo, fotos, leyenda in secciones:
                if fotos:
                    doc.add_page_break()
                    doc.add_heading(titulo, level=1)
                    grid = doc.add_table(rows=0, cols=2)
                    for i, f in enumerate(fotos):
                        if i % 2 == 0: row = grid.add_row().cells
                        row[i % 2].paragraphs[0].add_run().add_picture(f, width=Inches(2.5))
                    
                    if leyenda:
                        p_ley = doc.add_paragraph()
                        r_ley = p_ley.add_run(f"\n{leyenda}")
                        r_ley.font.size = Pt(10)
                        r_ley.font.bold = True

            # 3. Firma
            doc.add_paragraph("\n\n")
            tbl_f = doc.add_table(rows=2, cols=1)
            tbl_f.style = 'Table Grid'
            tbl_f.rows[0].cells[0].text = "FIRMA:"
            tbl_f.rows[1].height = Inches(1.5)

            target = io.BytesIO()
            doc.save(target)
            st.success("✅ Acta generada correctamente.")
            st.download_button("💾 GUARDAR WORD", target.getvalue(), f"Acta_{edar}.docx", use_container_width=True)
